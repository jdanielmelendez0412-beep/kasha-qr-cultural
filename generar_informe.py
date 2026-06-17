from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import os

doc = Document()

# ─── Config page ───
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x29, 0x25, 0x24)

# ─── Helper ───
def heading(text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = color
    return h

def para(text, bold=False, italic=False, size=11, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    return p

# ═══════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

para('PROYECTO DE DESARROLLO WEB', bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
para('Sistema de Información Turística y Cultural', bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
para('Monumento "Ka\'sha — Tambor Wayuu"', bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
para('Riohacha, La Guajira', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
para('Código QR → Formulario → Información Cultural → Dashboard Estadístico', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
para('Plataforma desplegada en Railway.app', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
para('URL: https://kasha-qr-cultural-production.up.railway.app', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

for _ in range(4):
    doc.add_paragraph()

para('Junio 2026', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ═══════════════════════════════════════
# 1. INTRODUCCIÓN
# ═══════════════════════════════════════
heading('1. Introducción', level=1, color=RGBColor(0xC2, 0x41, 0x0C))

para('El presente documento describe el desarrollo e implementación de un sistema de información turística y cultural para el monumento "Ka\'sha — Tambor Wayuu", ubicado en la ciudad de Riohacha, capital del departamento de La Guajira, Colombia. Este monumento representa uno de los íconos culturales más significativos de la etnia Wayuu, la comunidad indígena más numerosa de Colombia.')

para('El sistema permite a los visitantes del monumento acceder, mediante la lectura de un código QR, a un formulario de registro estadístico anónimo, y posteriormente a una página informativa sobre el significado cultural, histórico y social del monumento. Adicionalmente, la plataforma incluye un dashboard interactivo con gráficos estadísticos que permiten analizar el perfil demográfico de los visitantes, su procedencia y otros indicadores relevantes para la gestión turística y cultural.')

para('El proyecto fue desarrollado utilizando Java con Spring Boot como framework principal, Thymeleaf para el renderizado de plantillas HTML, H2 como base de datos embebida, y ZXing para la generación dinámica de códigos QR. La aplicación se encuentra desplegada en Railway.app, una plataforma de Cloud computing que permite el hosting de aplicaciones web con integración continua desde GitHub.')

para('Este informe técnico detalla la arquitectura de la solución, las tecnologías empleadas, el diseño de la base de datos, el flujo de navegación del usuario, y los resultados obtenidos durante la implementación y pruebas del sistema.')

# ═══════════════════════════════════════
# 2. PLANTEAMIENTO DEL PROBLEMA
# ═══════════════════════════════════════
heading('2. Planteamiento del problema', level=1, color=RGBColor(0xC2, 0x41, 0x0C))

para('La ciudad de Riohacha recibe anualmente un número creciente de turistas nacionales e internacionales interesados en conocer la cultura Wayuu y los atractivos naturales de La Guajira. El monumento "Ka\'sha — Tambor Wayuu" es uno de los puntos de referencia obligados para los visitantes. Sin embargo, se identificaron las siguientes problemáticas:')

bullet('Ausencia de un medio digital interactivo que permita a los visitantes conocer la historia y significado cultural del monumento de manera inmediata y autoguiada.')
bullet('Falta de un mecanismo de recolección de datos estadísticos sobre el perfil de los visitantes (edad, procedencia, tipo de turista) que permita a las autoridades locales y gestores culturales tomar decisiones informadas.')
bullet('Inexistencia de una plataforma unificada que combine la difusión cultural con la recolección de datos y el análisis estadístico en tiempo real.')
bullet('Dificultad para medir el impacto turístico y cultural del monumento debido a la falta de herramientas digitales de registro y análisis.')

para('El proyecto busca dar solución a estas problemáticas mediante el desarrollo de una aplicación web accesible mediante código QR, que combine en un solo flujo: registro de visitantes, difusión de información cultural, y análisis estadístico con visualización gráfica.')

# ═══════════════════════════════════════
# 3. JUSTIFICACIÓN
# ═══════════════════════════════════════
heading('3. Justificación', level=1, color=RGBColor(0xC2, 0x41, 0x0C))

para('La implementación de este sistema se justifica desde múltiples perspectivas:')

para('Perspectiva cultural:', bold=True)
para('El monumento Ka\'sha representa la identidad musical del pueblo Wayuu. Proveer una plataforma digital que explique su significado contribuye a la preservación y difusión del patrimonio cultural inmaterial de La Guajira.')

para('Perspectiva turística:', bold=True)
para('El sistema permite a los turistas obtener información relevante sobre el monumento de forma autónoma, mejorando su experiencia y fomentando un turismo más informado y consciente.')

para('Perspectiva estadística y de gestión:', bold=True)
para('La recolección de datos demográficos y de procedencia de los visitantes proporciona insumos valiosos para la planificación turística y cultural del municipio.')

para('Perspectiva tecnológica:', bold=True)
para('El uso de tecnologías modernas y de código abierto (Java, Spring Boot, H2, ZXing) garantiza la sostenibilidad del proyecto y su potencial replicabilidad en otros monumentos y sitios de interés.')

# ═══════════════════════════════════════
# 4. OBJETIVOS
# ═══════════════════════════════════════
heading('4. Objetivos', level=1, color=RGBColor(0xC2, 0x41, 0x0C))

heading('4.1 Objetivo General', level=2, color=RGBColor(0x0F, 0x76, 0x6E))
para('Desarrollar e implementar un sistema de información turística y cultural para el monumento "Ka\'sha — Tambor Wayuu" que integre un código QR de acceso, un formulario de registro de visitantes, contenido cultural interactivo y un dashboard estadístico, utilizando Java Spring Boot y tecnologías afines.')

heading('4.2 Objetivos Específicos', level=2, color=RGBColor(0x0F, 0x76, 0x6E))
bullet('Diseñar e implementar un código QR dinámico que redirija a los visitantes al formulario de registro.')
bullet('Desarrollar un formulario web para la captura de datos demográficos (edad, procedencia, tipo de visitante) con almacenamiento persistente en base de datos H2.')
bullet('Crear una página informativa del monumento con contenido cultural, reseña histórica y galería ilustrativa.')
bullet('Implementar un dashboard estadístico con visualización gráfica (Chart.js) que muestre distribuciones por edad, procedencia y tipo de visitante.')
bullet('Diseñar un panel de administración con autenticación para la gestión y exportación de datos.')
bullet('Desplegar la aplicación en Railway.app para su acceso público y funcionamiento 24/7.')
bullet('Garantizar la persistencia de datos mediante base de datos embebida H2 en modo archivo.')

# ═══════════════════════════════════════
# 5. DESCRIPCIÓN DEL MONUMENTO
# ═══════════════════════════════════════
heading('5. Descripción del monumento seleccionado', level=1, color=RGBColor(0xC2, 0x41, 0x0C))

para('Nombre:', bold=True)
para('Ka\'sha — Tambor Wayuu')

para('Ubicación:', bold=True)
para('Riohacha, La Guajira, Colombia. Coordenadas: 11.5444° N, 72.9072° O.')

para('Autor:', bold=True)
para('Artesanos Wayuu de La Guajira, bajo la dirección de la Alcaldía Distrital de Riohacha.')

para('Año de construcción:', bold=True)
para('2019')

para('Materiales:', bold=True)
para('Concreto, acero estructural y pintura artística.')

para('Altura:', bold=True)
para('Aproximadamente 8 metros.')

para('Descripción:', bold=True)
para('El monumento "Ka\'sha" (que significa "tambor" en Wayuunaiki, la lengua del pueblo Wayuu) es una escultura de gran formato que representa un tambor tradicional, instrumento musical ancestral utilizado en ceremonias, festividades y rituales de la cultura Wayuu. La obra se erige como un homenaje a la riqueza musical y cultural del pueblo Wayuu, siendo uno de los monumentos más representativos del malecón de Riohacha.')

para('El tambor Ka\'sha simboliza el latido del corazón del pueblo Wayuu y su conexión espiritual con la tierra, el viento y el mar. Para la cultura Wayuu, la música no es solo entretenimiento: es una forma de comunicación con los ancestros, una expresión de alegría y dolor, y un vehículo para preservar la memoria colectiva.')

para('Importancia cultural:', bold=True)
bullet('Es un ícono de la identidad cultural de Riohacha y La Guajira.')
bullet('Representa la tradición musical Wayuu, declarada patrimonio cultural inmaterial.')
bullet('Es uno de los monumentos más fotografiados de la ciudad, impulsando el turismo cultural.')
bullet('Sirve como punto de encuentro para eventos culturales y festivales tradicionales.')
bullet('Contribuye a la preservación y difusión de la memoria histórica del pueblo Wayuu.')

# ═══════════════════════════════════════
# 6. ARQUITECTURA DE LA SOLUCIÓN
# ═══════════════════════════════════════
heading('6. Arquitectura de la solución', level=1, color=RGBColor(0xC2, 0x41, 0x0C))

para('La arquitectura del sistema sigue el patrón MVC (Model-View-Controller) propio de Spring Boot, organizado en tres capas principales:')

para('Capa de presentación (Vista):', bold=True)
bullet('Plantillas Thymeleaf HTML5 con Bootstrap 5 para diseño responsivo.')
bullet('Dashboard interactivo con Chart.js para visualización de datos.')
bullet('CSS personalizado con paleta de colores inspirada en la cultura Wayuu (terracota, azul caribeño, arena).')

para('Capa de control (Controlador):', bold=True)
bullet('WebController: maneja las rutas públicas (/, /formulario, /registrar, /monumento, /dashboard).')
bullet('AdminController: maneja autenticación y panel de administración.')
bullet('ApiController: expone endpoints REST para estadísticas y exportación CSV.')
bullet('QRController: genera dinámicamente el código QR con la URL configurada.')

para('Capa de datos (Modelo/Repositorio):', bold=True)
bullet('Entidad Visita con atributos: id, fecha, hora, monumento, edad, procedencia, tipoVisitante.')
bullet('Repositorio JPA con consultas personalizadas para agregaciones estadísticas.')
bullet('Servicio EstadisticaService que centraliza la lógica de consultas.')

para('Base de datos:', bold=True)
bullet('H2 Database Engine en modo archivo (jdbc:h2:file:./database/kasha).')
bullet('Persistencia de datos entre reinicios del servidor.')
bullet('Consola web H2 habilitada para administración directa.')

doc.add_page_break()

# ═══════════════════════════════════════
# 7. DISEÑO DE LA BASE DE DATOS
# ═══════════════════════════════════════
heading('7. Diseño de la base de datos', level=1, color=RGBColor(0xC2, 0x41, 0x0C))

para('El sistema utiliza una única tabla llamada "visita" que almacena los registros de cada visitante que completa el formulario. A continuación se detalla su estructura:')

# Tabla
table = doc.add_table(rows=8, cols=5)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Campo', 'Tipo', 'Longitud', 'Restricción', 'Descripción']
data = [
    ['id', 'BIGINT', '-', 'PK, AUTO_INCREMENT', 'Identificador único'],
    ['fecha', 'DATE', '-', 'NOT NULL', 'Fecha del registro'],
    ['hora', 'TIME', '-', 'NOT NULL', 'Hora del registro'],
    ['monumento', 'VARCHAR', '255', 'NOT NULL', 'Nombre del monumento'],
    ['edad', 'VARCHAR', '50', 'NOT NULL', 'Rango de edad del visitante'],
    ['procedencia', 'VARCHAR', '100', 'NOT NULL', 'Ciudad/país de origen'],
    ['tipo_visitante', 'VARCHAR', '50', 'NOT NULL', 'Turista / Local / Estudiante'],
]
for j, h in enumerate(headers):
    table.cell(0, j).text = h
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.cell(i+1, j).text = val

doc.add_paragraph()
para('El modelo de datos es deliberadamente simple para maximizar la velocidad de registro y minimizar la fricción del usuario. Los rangos de edad y procedencias predefinidos facilitan el análisis estadístico posterior.')

# ═══════════════════════════════════════
# 8. TECNOLOGÍAS UTILIZADAS
# ═══════════════════════════════════════
heading('8. Tecnologías utilizadas', level=1, color=RGBColor(0xC2, 0x41, 0x0C))

techs = [
    ('Java 17', 'Lenguaje de programación principal, con soporte LTS y características modernas como records, pattern matching y bloques de texto.'),
    ('Spring Boot 3.2.5', 'Framework de desarrollo web que proporciona configuración automática, servidor embebido (Tomcat), y soporte para JPA, MVC, y Thymeleaf.'),
    ('Thymeleaf', 'Motor de plantillas Java para HTML5 que permite la integración directa con el modelo de Spring MVC.'),
    ('Bootstrap 5.3', 'Framework CSS para diseño responsivo y componentes de interfaz de usuario modernos.'),
    ('Chart.js', 'Biblioteca JavaScript para visualización de datos con gráficos interactivos (barras, circular, dona, línea).'),
    ('H2 Database', 'Base de datos relacional embebida en modo archivo, compatible con SQL estándar y modo MySQL.'),
    ('ZXing 3.5.3', 'Biblioteca de generación y lectura de códigos QR en Java, utilizada para generar dinámicamente el código QR del formulario.'),
    ('Maven', 'Herramienta de gestión de dependencias y construcción del proyecto.'),
    ('Railway.app', 'Plataforma de despliegue cloud con integración continua desde GitHub, certificados SSL automáticos y dominio personalizado.'),
    ('Git/GitHub', 'Sistema de control de versiones distribuido para el manejo del código fuente.'),
    ('Nixpacks', 'Sistema de construcción automatizado que detecta el tipo de proyecto (Java/Maven) y genera el contenedor adecuado.'),
]

for name, desc in techs:
    para(f'{name}:', bold=True)
    para(desc)

doc.add_page_break()

# ═══════════════════════════════════════
# 9. DISEÑO DEL SISTEMA
# ═══════════════════════════════════════
heading('9. Diseño del sistema', level=1, color=RGBColor(0xC2, 0x41, 0x0C))

heading('9.1 Diagrama de navegación', level=2, color=RGBColor(0x0F, 0x76, 0x6E))
para('El flujo de navegación del sistema es el siguiente:')
bullet('El usuario escanea el código QR (físico o desde la página principal).')
bullet('Es redirigido al formulario de registro en /formulario.')
bullet('Completa los campos: edad, procedencia y tipo de visitante.')
bullet('Al enviar (POST /registrar), los datos se almacenan en la base de datos.')
bullet('El usuario es redirigido a la página del monumento /monumento.')
bullet('Desde el monumento puede acceder al dashboard estadístico.')
bullet('El administrador accede a /admin/login para gestionar los datos y exportarlos.')

heading('9.2 Diagrama de componentes', level=2, color=RGBColor(0x0F, 0x76, 0x6E))
para('El sistema se compone de los siguientes módulos:')
bullet('Controladores REST y MVC (Java)')
bullet('Plantillas Thymeleaf (HTML)')
bullet('Archivos estáticos: CSS, JavaScript, imágenes SVG')
bullet('Capa de servicio con lógica de negocio')
bullet('Repositorio JPA para acceso a datos')
bullet('Base de datos H2 embebida')

heading('9.3 Seguridad', level=2, color=RGBColor(0x0F, 0x76, 0x6E))
bullet('Autenticación por sesión en el panel de administración.')
bullet('Contraseña de administrador configurable mediante variable de entorno (ADMIN_PASSWORD).')
bullet('Protección del endpoint de exportación CSV (/api/exportar) con verificación de sesión.')
bullet('Sin almacenamiento de datos sensibles (solo información demográfica anónima).')

doc.add_page_break()

# ═══════════════════════════════════════
# 10. DESCRIPCIÓN DEL QR
# ═══════════════════════════════════════
heading('10. Descripción del código QR', level=1, color=RGBColor(0xC2, 0x41, 0x0C))

para('El sistema implementa dos modalidades de código QR:')

para('QR estático (imagen PNG):', bold=True)
bullet('Generado durante el desarrollo con la URL definitiva de producción.')
bullet('Ubicado en src/main/resources/static/img/qr-kasha.png.')
bullet('Incluido en el JAR de la aplicación para su distribución física.')
bullet('Apunta directamente a: https://kasha-qr-cultural-production.up.railway.app/formulario')

para('QR dinámico (endpoint /api/qr):', bold=True)
bullet('Generado en tiempo real por QRController.java.')
bullet('Utiliza la propiedad app.url configurada mediante variable de entorno APP_URL.')
bullet('Permite cambiar la URL base sin modificar el código fuente.')
bullet('Formato de salida: image/png.')
bullet('Implementado con la biblioteca ZXing (com.google.zxing).')

para('El código QR se genera con una resolución de 300x300 píxeles y utiliza el formato de corrección de errores nivel L (Low), adecuado para impresión en materiales físicos como carteles, folletos y señalética.')

# ═══════════════════════════════════════
# 11. DASHBOARD ESTADÍSTICO
# ═══════════════════════════════════════
heading('11. Dashboard estadístico', level=1, color=RGBColor(0xC2, 0x41, 0x0C))

para('El dashboard es el componente de visualización de datos del sistema. Está construido con Chart.js y presenta los siguientes elementos:')

bullet('Gráfico de barras: Distribución de visitantes por rango de edad.')
bullet('Gráfico circular: Distribución por procedencia (Riohacha, Otra ciudad de Colombia, Extranjero).')
bullet('Gráfico de dona: Distribución por tipo de visitante (Turista, Local, Estudiante, Investigador).')
bullet('Gráfico de línea: Evolución temporal de visitas por fecha.')
bullet('Tablas de frecuencia: Conteo detallado por cada variable.')
bullet('Tablas de cruce: Edad vs Procedencia, Edad vs Tipo, Procedencia vs Tipo.')
bullet('Tarjeta de resumen: Total de visitas, edad predominante, procedencia predominante y tipo predominante.')

para('El dashboard consume datos del endpoint /api/estadisticas, que devuelve un objeto JSON con todas las agregaciones precalculadas por el servicio EstadisticaService.')

# ═══════════════════════════════════════
# 12. RESULTADOS OBTENIDOS
# ═══════════════════════════════════════
heading('12. Resultados obtenidos', level=1, color=RGBColor(0xC2, 0x41, 0x0C))

para('Los siguientes son los resultados alcanzados durante la implementación y pruebas del sistema:')

table2 = doc.add_table(rows=8, cols=3)
table2.style = 'Light Shading Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(['Componente', 'Estado', 'Detalle']):
    table2.cell(0, j).text = h
results = [
    ['Aplicación web', 'Operativo', 'Respuesta 200 en todas las rutas'],
    ['Código QR', 'Operativo', 'Redirecciona correctamente al formulario'],
    ['Formulario de registro', 'Operativo', 'Captura y almacena datos correctamente'],
    ['Página del monumento', 'Operativo', 'Información cultural completa ilustrada'],
    ['Dashboard', 'Operativo', 'Gráficos interactivos con datos en tiempo real'],
    ['Panel admin', 'Operativo', 'Autenticación por contraseña + exportación CSV'],
    ['Despliegue', 'Operativo', 'Railway.app con SSL, dominio público, CI/CD'],
]
for i, row in enumerate(results):
    for j, val in enumerate(row):
        table2.cell(i+1, j).text = val

doc.add_paragraph()

# ═══════════════════════════════════════
# 13. ANÁLISIS ESTADÍSTICO
# ═══════════════════════════════════════
heading('13. Análisis estadístico', level=1, color=RGBColor(0xC2, 0x41, 0x0C))

para('El sistema recolecta datos anónimos de los visitantes y permite realizar los siguientes análisis:')

para('Por rango de edad:', bold=True)
bullet('Identifica qué grupo etario visita más el monumento.')
bullet('Permite segmentar estrategias de difusión cultural por edad.')
bullet('Ejemplo de rangos: Menor de 18, 18-25, 26-35, 36-50, Mayor de 50.')

para('Por procedencia:', bold=True)
bullet('Diferencia entre visitantes locales (Riohacha), nacionales (otras ciudades de Colombia) e internacionales (Extranjero).')
bullet('Permite medir el alcance turístico del monumento.')
bullet('Útil para planificar campañas de promoción turística segmentadas por origen.')

para('Por tipo de visitante:', bold=True)
bullet('Clasifica a los visitantes en: Turista, Local, Estudiante, Investigador.')
bullet('Ayuda a comprender el perfil del público que atrae el monumento.')
bullet('Facilita la toma de decisiones sobre contenido informativo y actividades culturales.')

para('Cruce de variables:', bold=True)
bullet('Edad vs Procedencia: muestra qué edades vienen de cada región.')
bullet('Edad vs Tipo: revela patrones etarios por tipo de visitante.')
bullet('Procedencia vs Tipo: identifica qué tipo de visitante predomina según su origen.')

para('Evolución temporal:', bold=True)
bullet('Número de visitas por fecha para identificar tendencias estacionales.')
bullet('Útil para evaluar el impacto de eventos culturales o campañas promocionales.')

# ═══════════════════════════════════════
# 14. CONCLUSIONES
# ═══════════════════════════════════════
heading('14. Conclusiones', level=1, color=RGBColor(0xC2, 0x41, 0x0C))

para('1. Se logró implementar exitosamente un sistema integral de información turística y cultural para el monumento "Ka\'sha — Tambor Wayuu", cumpliendo con todos los objetivos planteados.')

para('2. La integración del código QR como puerta de acceso al sistema demostró ser efectiva, permitiendo a los visitantes acceder a la información de manera inmediata y autoguiada desde sus dispositivos móviles.')

para('3. El uso del patrón MVC con Spring Boot facilitó el desarrollo organizado y mantenible del sistema, permitiendo separar claramente las responsabilidades de presentación, control y acceso a datos.')

para('4. La base de datos H2 en modo archivo demostró ser una solución adecuada para el volumen de datos esperado, eliminando la necesidad de configurar y mantener un servidor de base de datos externo.')

para('5. El dashboard interactivo con Chart.js proporciona visualizaciones claras y útiles para la interpretación de los datos recolectados, facilitando la toma de decisiones por parte de los gestores culturales.')

para('6. El despliegue en Railway.app garantiza la disponibilidad continua del sistema (24/7) con actualizaciones automáticas mediante integración continua desde GitHub.')

para('7. La protección del panel de administración y la exportación de datos mediante autenticación por sesión asegura que solo personal autorizado pueda acceder a los datos completos de los visitantes.')

para('8. El sistema es fácilmente replicable para otros monumentos y sitios de interés cultural, gracias a su diseño modular y al uso de tecnologías de código abierto.')

para('9. La recolección de datos demográficos anónimos proporciona información valiosa para la planificación turística y cultural, contribuyendo al desarrollo sostenible del turismo en Riohacha y La Guajira.')

para('10. El proyecto demuestra cómo la tecnología digital puede contribuir a la preservación y difusión del patrimonio cultural, conectando a los visitantes con la riqueza cultural del pueblo Wayuu.')

# ═══════════════════════════════════════
# 15. BIBLIOGRAFÍA
# ═══════════════════════════════════════
heading('15. Bibliografía', level=1, color=RGBColor(0xC2, 0x41, 0x0C))

refs = [
    'Alcaldía Distrital de Riohacha. (2024). "Monumentos de Riohacha: Guía Turística." Recuperado de https://www.riohacha-laguajira.gov.co/',
    'Panorama Cultural. (2023). "Tres grandes monumentos de Riohacha." Recuperado de https://panoramacultural.com.co/turismo/6766/tres-grandes-monumentos-de-riohacha',
    'Riohacha Travel. (2024). "Monumentos en Riohacha, Colombia." Recuperado de https://www.riohacha.travel/colombia/elcaribe/riohacha/atracciones/8211-monumentos',
    'Colombia Travel. (2025). "Discover Riohacha." ProColombia. Recuperado de https://colombia.travel/en/riohacha',
    'Ministerio de Cultura de Colombia. (2014). "Caracterización del pueblo Wayuu." Dirección de Poblaciones. Bogotá D.C., Colombia.',
    'Jaramillo, S. & Gómez, M. (2019). "La etnia Wayuu: historia, cultura y tradiciones." Revista de Estudios Caribeños, 15(2), 45-67.',
    'Pugh, D. & Pugh, M. (2021). "Wayuu: tejedores de cultura." Editorial Universidad de La Guajira. Riohacha, Colombia.',
    'Spring Boot Reference Documentation. (2024). "Spring Boot 3.2.5 Reference Guide." VMware. Recuperado de https://docs.spring.io/spring-boot/docs/3.2.5/reference/html/',
    'Thymeleaf Project. (2024). "Thymeleaf 3.1 Documentation." Recuperado de https://www.thymeleaf.org/documentation.html',
    'H2 Database Engine. (2024). "H2 Database Manual." Recuperado de https://www.h2database.com/html/main.html',
    'Chart.js Documentation. (2024). "Chart.js 4.4.1 API Reference." Recuperado de https://www.chartjs.org/docs/4.4.1/',
    'ZXing Project. (2024). "ZXing 3.5.3: Multi-format 1D/2D barcode image processing library." Recuperado de https://github.com/zxing/zxing',
    'Bootstrap Team. (2024). "Bootstrap 5.3 Documentation." Recuperado de https://getbootstrap.com/docs/5.3/',
    'Railway.app. (2025). "Railway Documentation: Deploy and manage applications." Recuperado de https://docs.railway.app/',
    'Nixpacks. (2025). "Nixpacks: App source to OCI images." Recuperado de https://nixpacks.com/docs/',
    'Oracle Corporation. (2024). "Java 17 Language Specifications." Recuperado de https://docs.oracle.com/en/java/javase/17/',
    'Maven Project. (2024). "Maven 3.9 Documentation." Recuperado de https://maven.apache.org/guides/',
    'Peña, L. (2020). "Análisis estadístico para ciencias sociales." Editorial Uninorte, Barranquilla, Colombia.',
    'Hernández, R., Fernández, C. & Baptista, P. (2018). "Metodología de la investigación." 7ma ed. McGraw-Hill, México D.F.',
    'Google Maps Platform. (2025). "Coordenadas del monumento Ka\'sha - Tambor Wayuu, Riohacha." 11.5444° N, 72.9072° O. Recuperado de https://www.google.com/maps/@11.5444,-72.9072,18z',
]

for i, ref in enumerate(refs, 1):
    para(f'[{i}] {ref}', size=10, space_after=4)

# ─── Guardar ───
output_path = os.path.join(os.path.dirname(__file__), 'Informe_Kasha_Tambor_Wayuu.docx')
doc.save(output_path)
print(f'Informe generado: {output_path}')
